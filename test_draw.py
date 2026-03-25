#!/usr/bin/env python3
"""Pygments token → Pillow 画布 PNG（行号 + 语法高亮）"""
import os, tempfile
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from pygments.lexers import get_lexer_by_name, TextLexer
from pygments.token import Token, _TokenType
import pygments

# ── 颜色配置（Monokai 风格）──
COLORS = {
    Token.Text:               (0xF8, 0xF8, 0xF2),
    Token.Keyword:            (0xF9, 0x26, 0x72),
    Token.Keyword.Constant:  (0xFD, 0xB5, 0x05),
    Token.Name.Builtin:       (0xA6, 0xE2, 0x2E),
    Token.Name.Function:      (0x66, 0xD9, 0xEF),
    Token.Name.Class:         (0x66, 0xD9, 0xEF),
    Token.String:             (0xE6, 0xDB, 0x74),
    Token.Number:             (0xAE, 0x81, 0xFF),
    Token.Comment:            (0x75, 0x76, 0x6C),
    Token.Operator:           (0xF8, 0xF8, 0xF2),
    Token.Punctuation:        (0xF8, 0xF8, 0xF2),
    Token.Name.Decorator:      (0xFD, 0xB5, 0x05),
    Token.Name.Variable:       (0xF8, 0xF8, 0xF2),
    Token.Generic.Inserted:    (0xA6, 0xE2, 0x2E),
    Token.Generic.Deleted:     (0xF9, 0x26, 0x72),
    'background':             (0x27, 0x29, 0x33),
    'line_number':            (0x75, 0x76, 0x6C),
    'line_number_bg':         (0x3B, 0x3B, 0x4A),
    'default':                (0xF8, 0xF8, 0xF2),
}

BG = COLORS['background']

def get_color(token_type):
    """从 token 类型获取颜色，向上回溯找最接近的已定义颜色"""
    while token_type:
        if token_type in COLORS:
            return COLORS[token_type]
        token_type = token_type.parent
    return COLORS['default']

def draw_code_block(code_text, language=None, font_size=12, line_height=18,
                    left_pad=10, line_num_width=45):
    """在 Pillow 画布上绘制带行号和语法高亮的代码块，返回 PNG bytes"""
    try:
        if language and language not in ('mermaid', ''):
            lexer = get_lexer_by_name(language)
        else:
            lexer = TextLexer()
    except Exception:
        lexer = TextLexer()

    # 逐行解析 tokens
    lines = code_text.split('\n')

    # 逐行获取 tokens
    tokens_by_line = []
    for line in lines:
        tokens_for_line = list(pygments.lex(line, lexer=lexer))
        tokens_by_line.append(tokens_for_line)

    # 计算尺寸
    try:
        font_path = 'C:/Windows/Fonts/cour.ttf'
        font = ImageFont.truetype(font_path, font_size)
    except Exception:
        font = ImageFont.load_default()

    # 估算每行宽度
    max_width = 0
    for line in lines:
        try:
            w = font.getbbox(line)[2] if hasattr(font, 'getbbox') else font.getsize(line)[0]
        except Exception:
            w = font.getsize(line)[0]
        max_width = max(max_width, w)

    total_width = left_pad * 2 + line_num_width + max_width + 10
    total_height = len(lines) * line_height + 4

    img = Image.new('RGB', (total_width, total_height), BG)
    draw = ImageDraw.Draw(img)

    # 绘制背景边框
    draw.rectangle([0, 0, total_width-1, total_height-1], outline=(0x50, 0x50, 0x60), width=1)

    for row_idx, tokens in enumerate(tokens_by_line):
        y = 2 + row_idx * line_height

        # 行号背景
        draw.rectangle([0, y, line_num_width, y + line_height],
                       fill=COLORS['line_number_bg'])

        # 行号文字
        line_num_str = str(row_idx + 1)
        draw.text((left_pad, y + 2), line_num_str,
                  font=font, fill=COLORS['line_number'])

        # 代码 tokens
        x = left_pad + line_num_width
        for token_type, text in tokens:
            color = get_color(token_type)
            try:
                draw.text((x, y + 2), text, font=font, fill=color)
                w = font.getbbox(text)[2] if hasattr(font, 'getbbox') else font.getsize(text)[0]
            except Exception:
                w = font.getsize(text)[0]
            x += w

    buf = BytesIO()
    img.save(buf, 'PNG')
    buf.seek(0)
    return buf.getvalue()


# 测试
code = 'def rag_pipeline(query):\n    docs = vector_db.search(query, top_k=10)\n    context = format_docs(docs)\n    response = llm.generate(context)\n    return response'

png_bytes = draw_code_block(code, language='python')
tmp_png = os.path.join(tempfile.gettempdir(), 'test_pygments_code.png')
with open(tmp_png, 'wb') as f:
    f.write(png_bytes)
print(f'PNG 生成成功: {tmp_png}, 大小: {len(png_bytes)} bytes')

# 插入 Word
from docx import Document
from docx.shared import Pt, Inches
doc = Document()
p = doc.add_paragraph()
run = p.add_run()
run.add_picture(tmp_png, width=Inches(6))
doc.save(os.path.join(tempfile.gettempdir(), 'test_pygments_word.docx'))
print('Word 测试文档已保存')
