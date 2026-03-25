#!/usr/bin/env python3
"""
Markdown to Word converter tool (支持内嵌HTML、Mermaid图表渲染、个性化定制)

个性化定制：
- 全文微软雅黑
- 正文10.5号
- 表格9号
- 标题动态字号（最末级12.5，每升一级+2）

Usage:
    python md2docx.py <input.md> <output.docx>
"""

import sys
import os
import re
import urllib.request
import urllib.parse
import base64
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from html.parser import HTMLParser

# 代码块语法高亮（支持 Pygments）
try:
    from code_block import add_code_block_to_doc
except ImportError:
    # 如果 code_block.py 不存在，回退到朴素版本
    def add_code_block_to_doc(doc, code_text, language=None):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)


# ============ 个性化定制配置 ============
FONT_NAME = '微软雅黑'           # 全文字体
BODY_FONT_SIZE = 10.5           # 正文字号
TABLE_FONT_SIZE = 9             # 表格字号
MIN_HEADING_SIZE = 12.5          # 最末级标题字号
HEADING_SIZE_STEP = 2            # 每级标题递增字号

# 表情符号映射表
EMOJI_MAP = {
    '👨‍👩‍👦': '[表情：一家人]',
    '👚': '[表情：衣服]',
    '🥗': '[表情：食物]',
    '🏠': '[表情：房子]',
    '🏡': '[表情：房子]',
    '🚴': '[表情：骑车]',
    '🚴‍♂️': '[表情：骑车]',
    '🧋': '[表情：奶茶]',
    '⛹': '[表情：篮球]',
    '⛹️': '[表情：篮球]',
    '🛌': '[表情：睡觉]',
    '🛀': '[表情：洗澡]',
    '💀': '[表情：骷髅]',
    '💞': '[表情：爱心]',
    '👨‍👩‍👦‍👦': '[表情：一家人]',
    '❤️': '[表情：爱心]',
    '🩷': '[表情：爱心]',
    '👨‍👩‍👧': '[表情：一家人]',
    '👨‍👩‍👧‍👦': '[表情：一家人]',
}


def replace_emojis(text):
    """替换文本中的表情符号为文字描述"""
    result = text
    for emoji, replacement in EMOJI_MAP.items():
        result = result.replace(emoji, replacement)
    return result


def calculate_heading_sizes(max_level):
    """
    根据最大标题级别计算各级标题字号
    例如：max_level=3 -> {1: 16.5, 2: 14.5, 3: 12.5}
    """
    sizes = {}
    for level in range(1, max_level + 1):
        # 最末级(min_level) = MIN_HEADING_SIZE
        # 每升一级 +HEADING_SIZE_STEP
        sizes[level] = MIN_HEADING_SIZE + (max_level - level) * HEADING_SIZE_STEP
    return sizes


def set_run_font(run, font_name=None, font_size=None, bold=None, italic=None, color=None):
    """设置run的字体、字号、样式"""
    if font_name:
        run.font.name = font_name
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_horizontal_rule(doc):
    """在文档中插入水平分割线（模拟 Word 的 --- + 回车效果）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_paragraph_font(para, font_name=None, font_size=None):
    """设置段落默认字体和字号"""
    if font_name or font_size:
        for run in para.runs:
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if font_size:
                run.font.size = Pt(font_size)


def render_mermaid_to_image(mermaid_code, output_path):
    """将Mermaid代码渲染为图片（本地 mmdc 渲染）"""
    try:
        import subprocess
        import tempfile
        import shutil

        mermaid_code = mermaid_code.strip()

        # 将 Mermaid 代码写入临时文件（mmdc 需要从文件读取）
        temp_dir = tempfile.gettempdir()
        mmd_counter = len([f for f in os.listdir(temp_dir) if f.startswith('mermaid_tmp_')])
        mmd_path = os.path.join(temp_dir, f'mermaid_tmp_{mmd_counter}.mmd')

        with open(mmd_path, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)

        # 调用本地 mmdc 渲染
        skill_dir = os.path.dirname(os.path.abspath(__file__))
        mmdc_cmd = os.path.join(skill_dir, 'node_modules', '.bin', 'mmdc.cmd')

        result = subprocess.run(
            [mmdc_cmd, '-i', mmd_path, '-o', output_path, '-b', 'white'],
            capture_output=True,
            timeout=60,
            cwd=skill_dir
        )

        # 清理临时文件
        try:
            os.remove(mmd_path)
        except Exception:
            pass

        if result.returncode == 0 and os.path.exists(output_path):
            return True
        else:
            print(f"Mermaid渲染失败: {result.stderr.decode('utf-8', errors='ignore')}")
            return False

    except subprocess.TimeoutExpired:
        print("Mermaid渲染超时")
        return False
    except Exception as e:
        print(f"Mermaid渲染失败: {e}")
        return False


class HTMLTextExtractor(HTMLParser):
    """提取HTML中的纯文本，保留部分格式信息"""
    
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.in_blockquote = False
        self.current_alignment = None
        self.current_bold = False
        self.current_italic = False
        self.tag_stack = []
        
    def handle_starttag(self, tag, attrs):
        self.tag_stack.append(tag)
        attrs_dict = dict(attrs)
        
        if tag == 'blockquote':
            self.in_blockquote = True
            self.text_parts.append(('start_quote', None))
        elif tag == 'strong' or tag == 'b':
            self.current_bold = True
        elif tag == 'em' or tag == 'i':
            self.current_italic = True
        elif tag == 'p':
            style = attrs_dict.get('style', '')
            align_match = re.search(r'text-align:\s*(\w+)', style)
            if align_match:
                align = align_match.group(1)
                if align == 'center':
                    self.current_alignment = 'center'
                elif align == 'right':
                    self.current_alignment = 'right'
        elif tag == 'h1':
            self.text_parts.append(('heading', 1))
        elif tag == 'h2':
            self.text_parts.append(('heading', 2))
        elif tag == 'h3':
            self.text_parts.append(('heading', 3))
        elif tag == 'br':
            self.text_parts.append(('newline', None))
        elif tag == 'div':
            align = attrs_dict.get('align', '')
            if align == 'center':
                self.current_alignment = 'center'
    
    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()
        
        if tag == 'blockquote':
            self.in_blockquote = False
            self.text_parts.append(('end_quote', None))
        elif tag == 'strong' or tag == 'b':
            self.current_bold = False
        elif tag == 'em' or tag == 'i':
            self.current_italic = False
        elif tag == 'p' or tag == 'div':
            if not self.in_blockquote:
                self.current_alignment = None
    
    def handle_data(self, data):
        data = replace_emojis(data)
        if data.strip():
            flags = []
            if self.current_bold:
                flags.append('bold')
            if self.current_italic:
                flags.append('italic')
            if self.current_alignment:
                flags.append(f'align:{self.current_alignment}')
            self.text_parts.append(('text', data, tuple(flags)))
        elif '\n' in data or data == '\n':
            self.text_parts.append(('newline', None))


def strip_html_tags(html_text):
    """移除HTML标签，提取纯文本，并替换表情符号"""
    clean = re.sub(r'<[^>]+>', '', html_text)
    clean = clean.replace('&nbsp;', ' ')
    clean = clean.replace('&lt;', '<')
    clean = clean.replace('&gt;', '>')
    clean = clean.replace('&amp;', '&')
    clean = clean.replace('&quot;', '"')
    clean = replace_emojis(clean)
    return clean


def extract_text_with_format(html_text):
    """从HTML中提取文本，保留格式信息"""
    extractor = HTMLTextExtractor()
    try:
        extractor.feed(html_text)
    except:
        pass
    return extractor.text_parts


def parse_table(lines):
    """解析Markdown表格，过滤分隔行（|---|---|）"""
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        # 跳过分隔行：每个单元格都只含有 - 和 : 和空格
        if all(re.match(r'^[-:\s]+$', cell) for cell in cells if cell):
            continue
        rows.append(cells)
    return rows


# 表格斑马纹颜色（纯灰白交替，不用蓝色）
TABLE_ROW_EVEN_COLOR = 'F2F2F2'   # 浅灰背景（偶数数据行）


def _set_cell_bg(cell, hex_color):
    """给表格单元格设置背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table_to_doc(doc, rows, font_size=None):
    """添加表格到Word文档：表头灰底，数据行白底（符合Markdown默认渲染风格）"""
    if not rows or len(rows) < 1:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        is_header = (i == 0)
        # Markdown表格默认：表头灰底，数据行白底
        if is_header:
            row_color = TABLE_ROW_EVEN_COLOR  # 表头灰底
        else:
            row_color = None  # 数据行白底（无背景）

        for j, cell_text in enumerate(row):
            cell_obj = table.cell(i, j)

            if row_color:
                _set_cell_bg(cell_obj, row_color)

            para = cell_obj.paragraphs[0]
            for run in para.runs:
                run.text = ''

            clean = strip_html_tags(cell_text)
            fs = font_size or TABLE_FONT_SIZE
            nodes = parse_inline(clean)
            render_inline_nodes(para, nodes, FONT_NAME, fs, bold=is_header)

            for run in para.runs:
                run.font.name = FONT_NAME
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                except Exception:
                    pass
                run.font.size = Pt(fs)
                run.font.color.rgb = RGBColor(0, 0, 0)


def add_html_content_to_doc(doc, html_content, body_size=BODY_FONT_SIZE):
    """处理HTML内容块并添加到文档"""
    parts = extract_text_with_format(html_content)
    
    if not parts:
        text = strip_html_tags(html_content)
        if text.strip():
            p = doc.add_paragraph()
            process_formatting(p, text.strip(), FONT_NAME, body_size)
        return
    
    current_para = None
    quote_mode = False
    # 同一 <p> 段落内累积文本，最后统一用 AST 渲染
    pending_runs = []   # list of (text, bold, italic, align)

    def flush_pending(para):
        """把 pending_runs 渲染到段落，每段用 AST 解析"""
        for text, bold, italic, align in pending_runs:
            if align == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            nodes = parse_inline(text)
            render_inline_nodes(para, nodes, FONT_NAME, body_size,
                                bold=bold, italic=italic)
        pending_runs.clear()

    for part in parts:
        part_type = part[0]
        
        if part_type == 'start_quote':
            if current_para:
                flush_pending(current_para)
            quote_mode = True
            current_para = doc.add_paragraph()
            current_para.style = 'Quote'

        elif part_type == 'end_quote':
            if current_para:
                flush_pending(current_para)
            quote_mode = False
            current_para = None

        elif part_type == 'heading':
            if current_para:
                flush_pending(current_para)
            level = part[1]
            current_para = doc.add_heading('', level=level - 1 if level > 0 else 0)

        elif part_type == 'text':
            text = part[1]
            flags = part[2] if len(part) > 2 else ()

            if current_para is None:
                if quote_mode:
                    current_para = doc.add_paragraph(style='Quote')
                else:
                    current_para = doc.add_paragraph()

            align_flag = [f for f in flags if f.startswith('align:')]
            align = align_flag[0].split(':')[1] if align_flag else None
            pending_runs.append((text, 'bold' in flags, 'italic' in flags, align))

        elif part_type == 'newline':
            if current_para:
                flush_pending(current_para)
                current_para.add_run('\n')

    if current_para:
        flush_pending(current_para)


def md_to_docx(input_file, output_file):
    """转换Markdown文件到Word文档"""
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    
    # 清零 Normal 样式的段前段后间距，防止每段后自动多一行
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)
    
    lines = md_content.strip().split('\n')
    input_dir = os.path.dirname(os.path.abspath(input_file))
    
    # 第一遍扫描：检测最大标题级别
    max_heading_level = 0
    for line in lines:
        heading_match = re.match(r'^(#{1,6})\s+', line)
        if heading_match:
            level = len(heading_match.group(1))
            max_heading_level = max(max_heading_level, level)
    
    # 计算各级标题字号
    heading_sizes = calculate_heading_sizes(max_heading_level) if max_heading_level > 0 else {1: 14.5}
    print(f"检测到 {max_heading_level} 级标题")
    print(f"标题字号: {heading_sizes}")
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 处理HTML标签开头的行
        if line.strip().startswith('<') and not line.strip().startswith('```'):
            html_block = []
            html_block.append(line)
            
            if line.strip().startswith('<blockquote') or line.strip().startswith('<div'):
                depth = 1
                i += 1
                while i < len(lines) and depth > 0:
                    html_block.append(lines[i])
                    if '</blockquote>' in lines[i]:
                        depth -= 1
                    if '<blockquote' in lines[i] and '</blockquote>' not in lines[i]:
                        depth += 1
                    if lines[i].strip().startswith('</div>') and html_block[-1].strip() == '</div>':
                        if '</div>' in '\n'.join(html_block[:-1]):
                            pass
                        depth -= 1
                    i += 1
                
                html_content = '\n'.join(html_block)
                add_html_content_to_doc(doc, html_content)
                continue
            
            elif line.strip().startswith('<h') or line.strip().startswith('<p') or line.strip().startswith('<div'):
                add_html_content_to_doc(doc, line)
                i += 1
                continue
        
        # 处理Markdown引用块
        if line.startswith('> ') and not line.strip().startswith('>&gt;'):
            quote_lines = []
            while i < len(lines) and (lines[i].startswith('> ') or lines[i].startswith('>')):
                q_line = lines[i]
                if q_line.startswith('> '):
                    q_line = q_line[2:]
                elif q_line.startswith('>'):
                    q_line = q_line[1:]
                quote_lines.append(q_line.strip())
                i += 1

            full_quote = '\n'.join(quote_lines)
            if full_quote.strip().startswith('<'):
                add_html_content_to_doc(doc, full_quote)
            else:
                # 多行引用：每行一个段落，全部用 Quote 样式，行内格式走 AST 解析
                for j, ql in enumerate(quote_lines):
                    p = doc.add_paragraph()
                    p.style = 'Quote'
                    if ql:
                        process_formatting(p, ql, FONT_NAME, BODY_FONT_SIZE)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        # 空行也保留引用段落
                        run = p.add_run('')
                        run.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        # 处理代码块（包括Mermaid）
        if line.strip().startswith('```'):
            code_language = line.strip()[3:].lower()
            
            code_content = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_content.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_content)
            
            if code_language == 'mermaid':
                mermaid_code = code_text
                temp_dir = tempfile.gettempdir()
                img_counter = len([f for f in os.listdir(temp_dir) if f.startswith('mermaid_')])
                img_path = os.path.join(temp_dir, f'mermaid_{img_counter}.png')
                
                print(f"正在渲染Mermaid图表...")
                if render_mermaid_to_image(mermaid_code, img_path):
                    try:
                        from PIL import Image as PILImage
                        img = PILImage.open(img_path)
                        orig_w, orig_h = img.size  # 单位：像素
                        img.close()
                        
                        # 页面可用尺寸（单位：英寸）
                        # A4 纸 8.27" x 11.69"，左右页边距约 1"，上下约 1"
                        MAX_WIDTH_INCHES = 6.0   # 最大可用宽度
                        MAX_HEIGHT_INCHES = 8.0  # 最大可用高度（留余量给段落间距）
                        
                        # 计算最终宽度：优先满足宽度限制，高度超限则按高度缩放
                        # 如果宽度超限 → 按宽度缩放（高度等比缩小）
                        # 如果高度超限 → 按高度缩放（宽度等比缩小）
                        # 96 DPI 是 Word 默认图片分辨率
                        final_width_inches = orig_w / 96  # 像素转英寸
                        
                        if final_width_inches > MAX_WIDTH_INCHES:
                            # 宽度超标，按宽度缩放
                            final_width_inches = MAX_WIDTH_INCHES
                            print(f"  → 宽度超标，已缩放至 {MAX_WIDTH_INCHES}\"")
                        elif orig_h / 96 > MAX_HEIGHT_INCHES:
                            # 高度超标，按高度缩放
                            final_width_inches = orig_w * (MAX_HEIGHT_INCHES / orig_h)
                            print(f"  → 高度超标，已缩放")
                        
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(final_width_inches))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"Mermaid图表渲染成功 (原始: {orig_w}x{orig_h}px)")
                    except ImportError:
                        # 没有 Pillow，回退到默认宽度
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"Mermaid图表渲染成功")
                    except Exception as e:
                        print(f"添加图片失败: {e}")
                        p = doc.add_paragraph()
                        p.add_run('[Mermaid图表]')
                else:
                    p = doc.add_paragraph()
                    p.add_run('[Mermaid图表渲染失败，原代码如下:]')
                    run = p.add_run(code_text)
                    set_run_font(run, 'Courier New', 9)
            else:
                # 使用 Pygments 语法高亮（带行号和彩色 token）
                add_code_block_to_doc(doc, code_text, language=code_language)
            
            i += 1
            continue
        
        # 处理表格
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows, TABLE_FONT_SIZE)
            continue
        
        # 处理水平分割线 ---
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = strip_html_tags(text)
            
            # 添加标题并设置字号
            p = doc.add_heading('', level=level-1 if level > 0 else 0)
            
            # 用 AST 渲染标题文字（支持粗体等行内格式）
            heading_size = heading_sizes.get(level, MIN_HEADING_SIZE)
            nodes = parse_inline(text)
            render_inline_nodes(p, nodes, FONT_NAME, heading_size, bold=True)

            # 强制覆盖 Word 内置样式颜色为黑色，并统一字体字号
            for run in p.runs:
                run.font.name = FONT_NAME
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                except Exception:
                    pass
                run.font.size = Pt(heading_size)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            i += 1
            continue
        
        # 处理列表
        if re.match(r'^[-*]\s+', line):
            list_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            process_formatting(p, list_text, FONT_NAME, BODY_FONT_SIZE)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        # 处理图片
        img_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            if not os.path.isabs(img_path):
                img_path = os.path.join(input_dir, img_path)
            
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"添加图片失败: {e}")
                    p = doc.add_paragraph()
                    run = p.add_run(f'[图片: {alt_text}]')
                    set_run_font(run, FONT_NAME, BODY_FONT_SIZE)
            else:
                print(f"图片不存在: {img_path}")
                p = doc.add_paragraph()
                run = p.add_run(f'[图片: {alt_text} ({img_path})]')
                set_run_font(run, FONT_NAME, BODY_FONT_SIZE)
            i += 1
            continue
        
        # 处理链接
        link_match = re.search(r'\[([^\]]+)\]\(([^\)]+)\)', line)
        if link_match:
            line = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1: \2', line)
        
        # 处理空行：跳过，不额外加段落（段间距由 Normal 样式控制）
        if not line.strip():
            i += 1
            continue
        
        # 处理普通段落（直接走 AST，保留 **粗体** 等行内格式）
        raw_line = strip_html_tags(line) if line.strip().startswith('<') else line
        if raw_line.strip():
            p = doc.add_paragraph()
            process_formatting(p, raw_line, FONT_NAME, BODY_FONT_SIZE)
            # 强制黑色，防止继承样式颜色
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        i += 1
    
    doc.save(output_file)
    print(f"转换完成: {input_file} -> {output_file}")


# ============================================================
# AST 行内元素解析器
# 节点类型: text / bold / italic / bold_italic / code / link
# ============================================================

class InlineNode:
    """行内 AST 节点"""
    def __init__(self, node_type, content, children=None, href=None):
        self.type = node_type       # 'text'|'bold'|'italic'|'bold_italic'|'code'|'link'
        self.content = content      # 叶子节点文本（非叶子为 None）
        self.children = children or []  # 子节点列表
        self.href = href            # link 节点的 URL


def parse_inline(text):
    """
    将行内 Markdown 文本解析为 AST 节点列表。
    支持：**bold**、*italic*、***bold+italic***、`code`、[link](url)
    核心思路：
      1. 扫描 token（标记符 / 普通文本）
      2. 用栈维护嵌套关系
      3. 遇到匹配的闭合标记符，弹出栈帧并生成节点
    """
    tokens = _tokenize_inline(text)
    nodes, _ = _parse_tokens(tokens, 0, end_marker=None)
    return nodes


def _tokenize_inline(text):
    """
    词法分析：将字符串切成 token 列表
    每个 token: (type, value)
      type: 'marker'（*** ** * `）| 'link_open'（[）| 'link_mid'（](）| 'link_close'（)）| 'text'
    """
    # 匹配模式：先匹配长 marker，再匹配短 marker，防止 ** 被切成两个 *
    pattern = re.compile(
        r'(\*{3}|\*{2}|\*|`|\[|\]\(|(?<=\]\().*?(?=\))|\))',
    )
    tokens = []
    pos = 0
    # 改用手动扫描来正确处理 [text](url) 结构
    i = 0
    while i < len(text):
        # 检测链接 [label](url)
        link_m = re.match(r'\[([^\[\]]*?)\]\(([^)]*?)\)', text[i:])
        if link_m:
            tokens.append(('link', link_m.group(1), link_m.group(2)))
            i += link_m.end()
            continue
        # 检测行内代码 `code`
        code_m = re.match(r'`([^`]+?)`', text[i:])
        if code_m:
            tokens.append(('code', code_m.group(1), None))
            i += code_m.end()
            continue
        # 检测 ***
        if text[i:i+3] == '***':
            tokens.append(('marker', '***', None))
            i += 3
            continue
        # 检测 **
        if text[i:i+2] == '**':
            tokens.append(('marker', '**', None))
            i += 2
            continue
        # 检测 *
        if text[i] == '*':
            tokens.append(('marker', '*', None))
            i += 1
            continue
        # 普通文本：累积到下一个特殊字符
        j = i + 1
        while j < len(text) and text[j] not in ('*', '`', '['):
            # 检测是否到链接开始
            if text[j] == '[':
                break
            j += 1
        tokens.append(('text', text[i:j], None))
        i = j
    return tokens


def _parse_tokens(tokens, start, end_marker):
    """
    递归解析 token 列表，返回 (节点列表, 消耗到的位置)
    end_marker: 期望的结束标记符，None 表示解析到末尾
    """
    nodes = []
    i = start
    text_buf = ''

    def flush_text():
        nonlocal text_buf
        if text_buf:
            nodes.append(InlineNode('text', text_buf))
            text_buf = ''

    while i < len(tokens):
        tok_type, tok_val, tok_extra = tokens[i]

        if tok_type == 'text':
            text_buf += tok_val
            i += 1
            continue

        if tok_type == 'code':
            flush_text()
            nodes.append(InlineNode('code', tok_val))
            i += 1
            continue

        if tok_type == 'link':
            flush_text()
            nodes.append(InlineNode('link', tok_val, href=tok_extra))
            i += 1
            continue

        if tok_type == 'marker':
            marker = tok_val

            # 如果是期待的结束标记，停止递归
            if marker == end_marker:
                flush_text()
                return nodes, i + 1

            # 尝试作为开启标记，找到对应的闭合位置
            close_idx = _find_close(tokens, i + 1, marker)

            if close_idx is None:
                # 找不到闭合，当作普通文本
                text_buf += marker
                i += 1
                continue

            # 找到闭合，递归解析内部
            flush_text()
            inner_nodes, _ = _parse_tokens(tokens, i + 1, end_marker=marker)

            if marker == '***':
                node = InlineNode('bold_italic', None, children=inner_nodes)
            elif marker == '**':
                node = InlineNode('bold', None, children=inner_nodes)
            elif marker == '*':
                node = InlineNode('italic', None, children=inner_nodes)
            else:
                node = InlineNode('text', marker)

            nodes.append(node)
            # 跳到闭合标记之后
            i = close_idx + 1
            continue

        i += 1

    flush_text()
    return nodes, i


def _find_close(tokens, start, marker):
    """在 tokens[start:] 中查找第一个与 marker 匹配的闭合 token"""
    depth = 0
    for i in range(start, len(tokens)):
        t, v, _ = tokens[i]
        if t == 'marker':
            if v == marker:
                if depth == 0:
                    return i
                depth -= 1
            elif v in ('***', '**', '*'):
                depth += 1
    return None


def render_inline_nodes(para, nodes, font_name, font_size, bold=False, italic=False):
    """
    将 AST 节点列表渲染到 Word 段落。
    bold/italic 参数用于父节点传递样式（嵌套）
    """
    for node in nodes:
        if node.type == 'text':
            if node.content:
                run = para.add_run(node.content)
                set_run_font(run, font_name, font_size, bold=bold, italic=italic)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif node.type == 'bold':
            render_inline_nodes(para, node.children, font_name, font_size,
                                bold=True, italic=italic)

        elif node.type == 'italic':
            render_inline_nodes(para, node.children, font_name, font_size,
                                bold=bold, italic=True)

        elif node.type == 'bold_italic':
            render_inline_nodes(para, node.children, font_name, font_size,
                                bold=True, italic=True)

        elif node.type == 'code':
            run = para.add_run(node.content)
            set_run_font(run, 'Courier New', font_size - 0.5)
            # 经典代码块效果：深灰背景 + 深灰文字
            run.font.color.rgb = RGBColor(0x4E, 0x4E, 0x4E)
            # 加背景色（Word shd 元素）
            rPr = run._r.get_or_add_rPr()
            shd_el = OxmlElement('w:shd')
            shd_el.set(qn('w:val'), 'clear')
            shd_el.set(qn('w:color'), 'auto')
            shd_el.set(qn('w:fill'), 'D9D9D9')    # 浅灰背景
            rPr.append(shd_el)

        elif node.type == 'link':
            # 链接：显示文字，不插超链接（Word超链接复杂，暂显文本）
            run = para.add_run(node.content)
            set_run_font(run, font_name, font_size, bold=bold, italic=italic)
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 0)


def process_formatting(para, text, font_name, font_size):
    """
    将行内 Markdown 文本通过 AST 解析后渲染到段落。
    替代原来的正则逐行匹配，处理嵌套和边界更健壮。
    """
    nodes = parse_inline(text)
    render_inline_nodes(para, nodes, font_name, font_size)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在 {input_file}")
        sys.exit(1)
    
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    md_to_docx(input_file, output_file)
